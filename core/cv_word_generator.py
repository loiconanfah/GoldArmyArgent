import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, Any

def get_theme_config(theme_id: str) -> Dict[str, Any]:
    """
    Retourne la configuration complète du thème (polices, couleurs, etc.) 
    pour garantir un rendu Word (.docx) haut de gamme et professionnel.
    """
    themes = {
        "goldarmy": {
            "primary": RGBColor(196, 154, 108),       # Bronze/Or satin de luxe
            "secondary": RGBColor(115, 115, 115),     # Gris neutre
            "divider_hex": "C49A6C",                  # Hex or satin
            "font_title": "Calibri Light",
            "font_body": "Calibri",
            "body_color": RGBColor(51, 65, 85),       # Slate 700
        },
        "minimaliste": {
            "primary": RGBColor(30, 41, 59),          # Slate 800
            "secondary": RGBColor(100, 116, 139),     # Slate 500
            "divider_hex": "1E293B",
            "font_title": "Calibri Light",
            "font_body": "Calibri",
            "body_color": RGBColor(51, 65, 85),
        },
        "executive": {
            "primary": RGBColor(26, 54, 93),          # Navy profond
            "secondary": RGBColor(100, 116, 139),     # Slate 500
            "divider_hex": "1A365D",
            "font_title": "Georgia",                  # Police serif premium
            "font_body": "Georgia",
            "body_color": RGBColor(31, 41, 55),       # Dark Gray 800
        },
        "creatif": {
            "primary": RGBColor(13, 148, 136),        # Teal soutenu
            "secondary": RGBColor(100, 110, 120),
            "divider_hex": "0D9488",
            "font_title": "Calibri Light",
            "font_body": "Calibri",
            "body_color": RGBColor(51, 65, 85),
        },
        "classique": {
            "primary": RGBColor(17, 24, 39),          # Noir pur 900
            "secondary": RGBColor(107, 114, 128),     # Gray 500
            "divider_hex": "111827",
            "font_title": "Times New Roman",
            "font_body": "Times New Roman",
            "body_color": RGBColor(31, 41, 55),
        },
        "neon_tech": {
            "primary": RGBColor(3, 105, 161),         # Bleu ciel profond
            "secondary": RGBColor(100, 116, 139),
            "divider_hex": "0369A1",
            "font_title": "Calibri Light",
            "font_body": "Calibri",
            "body_color": RGBColor(51, 65, 85),
        },
        "scandinave": {
            "primary": RGBColor(47, 93, 80),          # Vert sapin
            "secondary": RGBColor(120, 130, 125),
            "divider_hex": "2F5D50",
            "font_title": "Calibri Light",
            "font_body": "Calibri",
            "body_color": RGBColor(51, 65, 85),
        },
        "timeline": {
            "primary": RGBColor(194, 65, 12),         # Terracotta / Coral sombre
            "secondary": RGBColor(120, 110, 105),
            "divider_hex": "C2410C",
            "font_title": "Calibri Light",
            "font_body": "Calibri",
            "body_color": RGBColor(51, 65, 85),
        },
    }
    return themes.get(theme_id, themes["goldarmy"])

def add_p_border_bottom(paragraph, hex_color="CCCCCC"):
    """Ajoute une bordure inférieure élégante sous le paragraphe (ligne de séparation)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is not None:
        pBdr.remove(bottom)
        
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Épaisseur : 6/8 pt
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)

def add_section_heading(document: Document, text: str, config: Dict[str, Any]):
    """Ajoute un titre de section stylisé avec une bordure inférieure fine."""
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = config["font_title"]
    run.font.color.rgb = config["primary"]
    
    add_p_border_bottom(p, config["divider_hex"])

def add_bullet_item(document: Document, text: str, font_name: str, body_color: RGBColor, bullet_color: RGBColor):
    """Ajoute une puce avec un retrait suspendu personnalisé pour un alignement optimal."""
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.15
    
    # Symbole de puce coloré
    run_bullet = p.add_run("•\t")
    run_bullet.font.name = font_name
    run_bullet.font.size = Pt(9.5)
    run_bullet.font.color.rgb = bullet_color
    run_bullet.bold = True
    
    # Texte de la puce nettoyé des puces brutes
    clean_text = text.strip()
    if clean_text.startswith("•"):
        clean_text = clean_text[1:]
    elif clean_text.startswith("-"):
        clean_text = clean_text[1:]
    elif clean_text.startswith("–"):
        clean_text = clean_text[1:]
    clean_text = clean_text.strip()
    
    run_text = p.add_run(clean_text)
    run_text.font.name = font_name
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = body_color

def build_header(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any]):
    """Génère l'en-tête du CV (Nom, Titre, Coordonnées) de façon centrée et épurée."""
    # Nom du candidat
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    
    run_name = p_name.add_run(cv_data.get("full_name", "Candidat").upper())
    run_name.bold = True
    run_name.font.size = Pt(20)
    run_name.font.name = config["font_title"]
    run_name.font.color.rgb = config["primary"]
    
    # Titre professionnel
    title = cv_data.get("title", "")
    if title:
        p_title = document.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(6)
        
        run_title = p_title.add_run(title.upper())
        run_title.font.size = Pt(11)
        run_title.font.name = config["font_title"]
        run_title.font.color.rgb = config["secondary"]
        run_title.bold = True
        
    # Coordonnées (Email, Tél, LinkedIn, etc.)
    contact_parts = []
    if cv_data.get("email"): contact_parts.append(cv_data["email"])
    if cv_data.get("phone"): contact_parts.append(cv_data["phone"])
    if cv_data.get("location"): contact_parts.append(cv_data["location"])
    if cv_data.get("linkedin"): contact_parts.append(cv_data["linkedin"])
    if cv_data.get("github"): contact_parts.append(cv_data["github"])
    
    if contact_parts:
        p_contact = document.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_before = Pt(0)
        p_contact.paragraph_format.space_after = Pt(12)
        
        run_contact = p_contact.add_run("  •  ".join(contact_parts))
        run_contact.font.size = Pt(9.5)
        run_contact.font.name = config["font_body"]
        run_contact.font.color.rgb = config["secondary"]
        
        # Bordure sous le header
        add_p_border_bottom(p_contact, config["divider_hex"])

def add_summary(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any]):
    """Ajoute le résumé de profil."""
    summary = cv_data.get("summary", "")
    if summary:
        add_section_heading(document, "Profil Professionnel", config)
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(summary)
        run.font.size = Pt(9.5)
        run.font.name = config["font_body"]
        run.font.color.rgb = config["body_color"]

def add_experiences(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any], printable_width: Inches):
    """Ajoute l'historique des expériences avec alignement dynamique à droite des dates."""
    experiences = cv_data.get("experiences", [])
    if not experiences: return
    
    add_section_heading(document, "Expériences Professionnelles", config)
    
    for exp in experiences:
        if isinstance(exp, str):
            add_bullet_item(document, exp, config["font_body"], config["body_color"], config["primary"])
            continue
            
        p = document.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(printable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        
        # Intitulé de poste
        run_title = p.add_run(exp.get("title", ""))
        run_title.bold = True
        run_title.font.size = Pt(10.5)
        run_title.font.name = config["font_title"]
        run_title.font.color.rgb = RGBColor(17, 24, 39)
        
        # Entreprise & Localisation
        company = exp.get("company", "")
        loc = exp.get("location", "")
        co_str = ""
        if company:
            co_str += f" — {company}"
        if loc:
            co_str += f" ({loc})"
            
        if co_str:
            run_co = p.add_run(co_str)
            run_co.font.size = Pt(9.5)
            run_co.font.name = config["font_body"]
            run_co.font.color.rgb = config["secondary"]
            run_co.font.italic = True
            
        # Dates à droite
        start = exp.get("start_date", "")
        end = exp.get("end_date", "")
        date_str = f"{start} – {end}" if (start or end) else ""
        if date_str:
            p.add_run("\t")
            run_date = p.add_run(date_str)
            run_date.font.size = Pt(9.5)
            run_date.font.name = config["font_body"]
            run_date.font.color.rgb = config["primary"]
            run_date.bold = True
            
        # Puces de réalisation
        for bullet in exp.get("bullets", []):
            if bullet.strip():
                add_bullet_item(document, bullet, config["font_body"], config["body_color"], config["primary"])

def add_projects(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any], printable_width: Inches):
    """Ajoute les projets personnels avec description."""
    projects = cv_data.get("projects", [])
    if not projects: return
    
    add_section_heading(document, "Projets Personnels", config)
    
    for proj in projects:
        p = document.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(printable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        
        # Nom du projet
        run_name = p.add_run(proj.get("name", ""))
        run_name.bold = True
        run_name.font.size = Pt(10.5)
        run_name.font.name = config["font_title"]
        run_name.font.color.rgb = RGBColor(17, 24, 39)
        
        # Date du projet à droite
        date = proj.get("date") or proj.get("year") or proj.get("start_date")
        if date:
            p.add_run("\t")
            run_date = p.add_run(str(date))
            run_date.font.size = Pt(9.5)
            run_date.font.name = config["font_body"]
            run_date.font.color.rgb = config["primary"]
            run_date.bold = True
            
        # Description
        desc = proj.get("description", "")
        if desc:
            p_desc = document.add_paragraph()
            p_desc.paragraph_format.space_before = Pt(0)
            p_desc.paragraph_format.space_after = Pt(2)
            p_desc.paragraph_format.keep_with_next = True
            
            run_desc = p_desc.add_run(desc)
            run_desc.font.size = Pt(9.5)
            run_desc.font.name = config["font_body"]
            run_desc.font.italic = True
            run_desc.font.color.rgb = config["secondary"]
            
        # Puces
        for bullet in proj.get("bullets", []):
            if bullet.strip():
                add_bullet_item(document, bullet, config["font_body"], config["body_color"], config["primary"])

def add_education(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any], printable_width: Inches):
    """Ajoute la section éducation avec les dates à droite."""
    education = cv_data.get("education", [])
    if not education: return
    
    add_section_heading(document, "Formation", config)
    
    for edu in education:
        p = document.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(printable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        
        # Diplôme
        run_deg = p.add_run(edu.get("degree", ""))
        run_deg.bold = True
        run_deg.font.size = Pt(10)
        run_deg.font.name = config["font_title"]
        run_deg.font.color.rgb = RGBColor(17, 24, 39)
        
        # Institution
        inst = edu.get("institution") or edu.get("school", "")
        if inst:
            run_inst = p.add_run(f" — {inst}")
            run_inst.font.size = Pt(9.5)
            run_inst.font.name = config["font_body"]
            run_inst.font.color.rgb = config["secondary"]
            
        # Métadonnées à droite (Lieu & Année)
        meta_parts = []
        loc = edu.get("location", "")
        year = edu.get("year", "")
        if loc: meta_parts.append(loc)
        if year: meta_parts.append(str(year))
        
        if meta_parts:
            p.add_run("\t")
            run_meta = p.add_run(" · ".join(meta_parts))
            run_meta.font.size = Pt(9.5)
            run_meta.font.name = config["font_body"]
            run_meta.font.color.rgb = config["primary"]
            run_meta.bold = True

def add_skills(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any]):
    """Ajoute la grille de compétences."""
    skills = cv_data.get("skills", {})
    if not skills: return
    
    add_section_heading(document, "Compétences", config)
    
    if isinstance(skills, dict):
        for category, items in skills.items():
            if not items: continue
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.15)
            
            run_cat = p.add_run(f"{category} : ")
            run_cat.bold = True
            run_cat.font.size = Pt(9.5)
            run_cat.font.name = config["font_body"]
            run_cat.font.color.rgb = config["primary"]
            
            if isinstance(items, list):
                skills_str = ", ".join(str(i) for i in items)
            else:
                skills_str = str(items)
                
            run_items = p.add_run(skills_str)
            run_items.font.size = Pt(9.5)
            run_items.font.name = config["font_body"]
            run_items.font.color.rgb = config["body_color"]
    elif isinstance(skills, list):
         p = document.add_paragraph()
         p.paragraph_format.space_before = Pt(2)
         p.paragraph_format.space_after = Pt(3)
         p.paragraph_format.left_indent = Inches(0.15)
         
         run = p.add_run(", ".join(str(s) for s in skills))
         run.font.size = Pt(9.5)
         run.font.name = config["font_body"]
         run.font.color.rgb = config["body_color"]

def add_languages(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any]):
    """Ajoute les langues maîtrisées."""
    languages = cv_data.get("languages", [])
    if not languages: return
    
    add_section_heading(document, "Langues", config)
    
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.15)
    
    lang_strs = []
    for lang in languages:
        if isinstance(lang, str):
            lang_strs.append(lang)
        elif isinstance(lang, dict):
            l_name = lang.get("language", "")
            l_prof = lang.get("proficiency", "")
            if l_name:
                lang_strs.append(f"{l_name} ({l_prof})" if l_prof else l_name)
                
    run = p.add_run("  •  ".join(lang_strs))
    run.font.size = Pt(9.5)
    run.font.name = config["font_body"]
    run.font.color.rgb = config["body_color"]

def add_certifications(document: Document, cv_data: Dict[str, Any], config: Dict[str, Any]):
    """Ajoute les certifications obtenues."""
    certs = cv_data.get("certifications", [])
    if not certs: return
    
    add_section_heading(document, "Certifications", config)
    
    for cert in certs:
        if isinstance(cert, str):
            text = cert
        elif isinstance(cert, dict):
            c_name = cert.get("name", "")
            c_iss = cert.get("issuer", "")
            c_year = cert.get("year", "")
            text = c_name
            if c_iss: text += f" ({c_iss})"
            if c_year: text += f" - {c_year}"
            
        add_bullet_item(document, text, config["font_body"], config["body_color"], config["primary"])

def generate_cv_word(cv_data: Dict[str, Any], theme_id: str = "goldarmy") -> bytes:
    """
    Génère un CV au format Word (.docx) à partir des données JSON.
    Applique une mise en page d'excellence (Executive 1-column layout) 
    qui allie un design premium épuré et une compatibilité ATS absolue de 100 %.
    """
    document = Document()
    config = get_theme_config(theme_id)
    
    # Configuration des marges à 0.75" de chaque côté (compromis parfait)
    sections = document.sections
    for section in sections:
        section.page_width = Inches(8.5) # Format US Letter standardisé pour le calcul
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    printable_width = Inches(8.5) - Inches(0.75) - Inches(0.75) # Exactement 7.0 pouces imprimables
    
    # Définition du style et police de base
    style = document.styles['Normal']
    font = style.font
    font.name = config["font_body"]
    font.size = Pt(9.5)
    font.color.rgb = config["body_color"]
    
    # Construction séquentielle et ordonnée du CV
    build_header(document, cv_data, config)
    add_summary(document, cv_data, config)
    add_experiences(document, cv_data, config, printable_width)
    add_projects(document, cv_data, config, printable_width)
    add_education(document, cv_data, config, printable_width)
    add_skills(document, cv_data, config)
    add_languages(document, cv_data, config)
    add_certifications(document, cv_data, config)
    
    # Écriture dans le flux de mémoire
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.read()

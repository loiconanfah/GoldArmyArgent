"""
Service de génération de CV en format DOCX optimisé ATS.
Utilise python-docx pour créer un fichier Word propre et parsable par les ATS.
"""
import io
import json
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Ajoute un titre de section avec ligne horizontale."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    _set_font(run, bold=True, size=11, color=(31, 73, 125))
    
    # Ligne horizontale sous le titre
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_bullet(doc: Document, text: str):
    """Ajoute un bullet point propre (compatible ATS)."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(text.lstrip('•').lstrip('- ').strip())
    _set_font(run, size=10)
    return para


def generate_cv_docx(cv_data: Dict[str, Any]) -> bytes:
    """
    Génère un fichier DOCX ATS-optimisé à partir des données structurées du CV.
    Retourne les bytes du fichier.
    """
    doc = Document()

    # ── Marges serrées (ATS friendly) ──────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Style par défaut ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════
    # EN-TÊTE : Nom, titre, contact
    # ═══════════════════════════════════════════════════════════════════
    name = cv_data.get("full_name", "Candidat")
    title = cv_data.get("title", "")

    # Nom
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    run = name_para.add_run(name)
    _set_font(run, bold=True, size=18, color=(31, 73, 125))

    # Titre
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(4)
        run = title_para.add_run(title)
        _set_font(run, bold=False, size=12, color=(89, 89, 89))

    # Infos de contact sur une ligne
    contact_parts = []
    for field in ["email", "phone", "location", "linkedin", "github"]:
        val = cv_data.get(field, "")
        if val and val.lower() not in ["", "n/a", "null", "none"]:
            contact_parts.append(val)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(8)
        run = contact_para.add_run("  |  ".join(contact_parts))
        _set_font(run, size=9, color=(89, 89, 89))

    # ═══════════════════════════════════════════════════════════════════
    # RÉSUMÉ PROFESSIONNEL
    # ═══════════════════════════════════════════════════════════════════
    summary = cv_data.get("summary", "")
    if summary:
        _add_heading(doc, "Résumé Professionnel")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(summary)
        _set_font(run, size=10)

    # ═══════════════════════════════════════════════════════════════════
    # EXPÉRIENCES PROFESSIONNELLES
    # ═══════════════════════════════════════════════════════════════════
    experiences = cv_data.get("experiences", [])
    if experiences:
        _add_heading(doc, "Expériences Professionnelles")
        for exp in experiences:
            # Titre du poste | Entreprise
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(6)
            exp_para.paragraph_format.space_after = Pt(0)
            run = exp_para.add_run(exp.get("title", ""))
            _set_font(run, bold=True, size=11)
            
            company = exp.get("company", "")
            location = exp.get("location", "")
            if company:
                run2 = exp_para.add_run(f"  —  {company}")
                _set_font(run2, size=11, color=(89, 89, 89))
            
            # Dates | Lieu
            date_parts = []
            start = exp.get("start_date", "")
            end = exp.get("end_date", "Présent")
            if start:
                date_parts.append(f"{start} – {end}")
            if location:
                date_parts.append(location)
            
            if date_parts:
                date_para = doc.add_paragraph()
                date_para.paragraph_format.space_before = Pt(0)
                date_para.paragraph_format.space_after = Pt(2)
                run = date_para.add_run("  |  ".join(date_parts))
                _set_font(run, size=9, color=(128, 128, 128))

            # Bullets de réalisations
            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    _add_bullet(doc, bullet)

    # ═══════════════════════════════════════════════════════════════════
    # COMPÉTENCES TECHNIQUES
    # ═══════════════════════════════════════════════════════════════════
    skills = cv_data.get("skills", {})
    if skills:
        _add_heading(doc, "Compétences Techniques")
        for category, items in skills.items():
            if not items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run_label = p.add_run(f"{category}: ")
            _set_font(run_label, bold=True, size=10)
            run_vals = p.add_run(", ".join(items) if isinstance(items, list) else str(items))
            _set_font(run_vals, size=10)

    # ═══════════════════════════════════════════════════════════════════
    # FORMATION
    # ═══════════════════════════════════════════════════════════════════
    education = cv_data.get("education", [])
    if education:
        _add_heading(doc, "Formation")
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(4)
            edu_para.paragraph_format.space_after = Pt(0)
            run = edu_para.add_run(edu.get("degree", ""))
            _set_font(run, bold=True, size=10)
            
            inst = edu.get("institution", "")
            loc = edu.get("location", "")
            year = edu.get("year", "")
            
            details = []
            if inst:
                details.append(inst)
            if loc:
                details.append(loc)
            if year:
                details.append(year)
            
            if details:
                det_para = doc.add_paragraph()
                det_para.paragraph_format.space_before = Pt(0)
                det_para.paragraph_format.space_after = Pt(2)
                run = det_para.add_run("  |  ".join(details))
                _set_font(run, size=9, color=(89, 89, 89))

    # ═══════════════════════════════════════════════════════════════════
    # CERTIFICATIONS
    # ═══════════════════════════════════════════════════════════════════
    certifications = cv_data.get("certifications", [])
    if certifications:
        _add_heading(doc, "Certifications")
        for cert in certifications:
            if cert.strip():
                _add_bullet(doc, cert)

    # ═══════════════════════════════════════════════════════════════════
    # LANGUES
    # ═══════════════════════════════════════════════════════════════════
    languages = cv_data.get("languages", [])
    if languages:
        _add_heading(doc, "Langues")
        lang_para = doc.add_paragraph()
        lang_para.paragraph_format.space_after = Pt(4)
        run = lang_para.add_run("  |  ".join(languages))
        _set_font(run, size=10)

    # ── Sauvegarder en mémoire et retourner les bytes ──────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
